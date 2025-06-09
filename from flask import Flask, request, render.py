from flask import Flask, request, render_template, send_from_directory, url_for
import re
import os
from docx import Document

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []

    # Extract visible text from paragraphs with section numbers
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():  # Only add non-empty paragraphs
            full_text.append(f"{i}. {para.text}")
    
    # Extract text from tables
    for table_num, table in enumerate(doc.tables, 1):
        full_text.append(f"\nTable {table_num}:")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))
    
    # Extract text from headers and footers
    for section_num, section in enumerate(doc.sections, 1):
        # Extract headers
        header = section.header
        if any(paragraph.text.strip() for paragraph in header.paragraphs):
            full_text.append(f"\nHeader Section {section_num}:")
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
        
        # Extract footers
        footer = section.footer
        if any(paragraph.text.strip() for paragraph in footer.paragraphs):
            full_text.append(f"\nFooter Section {section_num}:")
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
    
    return '\n'.join(full_text)

def extract_codes_from_text(text, codes):
    codes_set = set(codes)
    found_codes = set()
    
    # Match each code in the text
    for code in codes_set:
        if code in text:
            found_codes.add(code)
    
    return found_codes

def extract_codes_and_descriptions(doc):
    codes_with_descriptions = {}
    
    # First pass: Look for codes in the Results section
    in_results_section = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Check if we're entering the Results section
        if "Results" in text:
            in_results_section = True
            continue
            
        if in_results_section and text:
            # Look for "-Top Racing Stripe" or similar patterns
            if text.startswith('-'):
                parts = text.split('CJTA')  # Split on common code prefix
                if len(parts) > 1:
                    desc = parts[0].replace('-', '').strip()
                    code = 'CJTA' + parts[1].strip()
                    codes_with_descriptions[code] = desc
                    
    # Second pass: Look in tables
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
                
            row_text = ' '.join(cell.text.strip() for cell in row.cells)
            
            # Look for "Over-the-Top Stripe Package (47B) CJTAB" pattern
            if 'Package' in row_text:
                match = re.search(r'(.*?Package.*?)\s+([A-Z0-9]{4,5})\s*$', row_text)
                if match:
                    desc, code = match.groups()
                    codes_with_descriptions[code] = desc.strip()
                    continue
            
            # Look for "Dark Horse Appearance Package (700A) (54C)" pattern
            if 'Dark Horse' in row_text:
                match = re.search(r'(Dark Horse.*?)\s+([A-Z0-9]{4,5})[/$]', row_text)
                if match:
                    desc, code = match.groups()
                    codes_with_descriptions[code] = desc.strip()
                    continue
                    
            # Look for code at end of cell with description before it
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    match = re.search(r'(.*?)\s+([A-Z0-9]{4,5})\s*$', cell_text)
                    if match:
                        desc, code = match.groups()
                        if len(desc) > 5:  # Only store if there's a meaningful description
                            codes_with_descriptions[code] = desc.strip()
    
    # Third pass: Look for any remaining codes in paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Look for "GT and GT Premium Over-the-Top Stripe Package (47K) CJTAK" pattern
            if 'Package' in text:
                match = re.search(r'(.*?Package.*?)\s+([A-Z0-9]{4,5})\s*$', text)
                if match:
                    desc, code = match.groups()
                    if code not in codes_with_descriptions:
                        codes_with_descriptions[code] = desc.strip()
                        
            # Look for any other code with description
            match = re.search(r'(.*?)\s+([A-Z0-9]{4,5})\s*$', text)
            if match:
                desc, code = match.groups()
                if len(desc) > 5 and code not in codes_with_descriptions:
                    codes_with_descriptions[code] = desc.strip()
                    
    return codes_with_descriptions

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files.get('file')
        file2 = request.files.get('file2')
        input_codes = request.form.get('input_codes', '')
        voci_codes_text = request.form.get('voci_codes', '')

        if file:
            # Load the document
            doc = Document(file)
            
            # Extract descriptions first
            codes_with_descriptions = extract_codes_and_descriptions(doc)
            
            # Convert DOCX to text for code matching
            text = extract_text_from_docx(file)
            
            # Parse input WERS codes from newline-separated input while preserving order
            input_codes_list = [code for code in re.findall(r'[A-Z0-9]{5}', input_codes)]
            
            # Parse VOCI codes from pasted text
            voci_codes_list = re.findall(r'[A-Z0-9]{5}', voci_codes_text)
            
            # Extract codes found in Word document 1
            found_codes_doc1 = extract_codes_from_text(text, input_codes_list)

            found_codes_doc2 = set()
            if file2:
                # Extract text from second document
                text2 = extract_text_from_docx(file2)
                found_codes_doc2 = extract_codes_from_text(text2, input_codes_list)
            
            # Determine which codes are from VOCI alone
            voci_alone_codes = set(voci_codes_list) - found_codes_doc1 - found_codes_doc2
            
            # Determine codes common to both VOCI and Word documents
            common_voci_and_doc1 = set(voci_codes_list).intersection(found_codes_doc1)
            common_voci_and_doc2 = set(voci_codes_list).intersection(found_codes_doc2)
            
            # Combine all found codes for display
            combined_codes = found_codes_doc1.union(found_codes_doc2)
            
            # Modify the code_results creation to include descriptions
            code_results = []
            for code in input_codes_list:
                if code in voci_alone_codes:
                    source = 'VOCI Only'
                elif code in common_voci_and_doc1 and code in common_voci_and_doc2:
                    source = 'Both VOCI and WERS Document 1 and 2'
                elif code in common_voci_and_doc1:
                    source = 'Both VOCI and WERS Document 1'
                elif code in common_voci_and_doc2:
                    source = 'Both VOCI and WERS Document 2'
                elif code in found_codes_doc1:
                    source = 'WERS Document 1 Only'
                elif code in found_codes_doc2:
                    source = 'WERS Document 2 Only'
                else:
                    continue
                
                result_dict = {
                    'source': source,
                    'description': codes_with_descriptions.get(code, '')
                }
                code_results.append((code, result_dict))
            
            # Add VOCI-only codes with their descriptions
            for code in voci_codes_list:
                if code not in input_codes_list:
                    result_dict = {
                        'source': 'VOCI Only',
                        'description': codes_with_descriptions.get(code, '')
                    }
                    code_results.append((code, result_dict))
            
            # Save text file with extracted text and results
            text_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'file.txt')
            
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write("WERS Code Analysis Results\n")
                f.write("=========================\n\n")
                
                # Write Document 1 content
                f.write("Extracted Text from Document 1:\n")
                f.write("--------------------------\n")
                f.write(text)
                f.write("\n\n")
                
                # Write Document 2 content if it exists
                if file2:
                    f.write("Extracted Text from Document 2:\n")
                    f.write("--------------------------\n")
                    f.write(text2)
                    f.write("\n\n")
                
                # Write analysis results
                f.write("Analysis Results:\n")
                f.write("----------------\n")
                for code, result in code_results:
                    f.write(f"{code}: {result['source']} - {result['description']}\n")
            
            return render_template('display.html', 
                                   code_results=code_results,
                                   file_txt_url=url_for('download_file', filename='file.txt'))

    return render_template('upload.html')

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@app.route('/process', methods=['POST'])
def process():
    # ... existing file handling code ...
    
    doc1_path = os.path.join(app.config['UPLOAD_FOLDER'], doc1.filename)
    doc1.save(doc1_path)
    
    # Extract codes and descriptions from the document
    codes_with_descriptions = extract_codes_and_descriptions(Document(doc1_path))
    
    # Process WERS codes as before
    results = []
    # ... existing code processing ...
    
    # Add descriptions to results
    for result in results:
        code = result['code']
        result['description'] = codes_with_descriptions.get(code, '')
    
    return render_template('display.html', results=results)

if __name__ == '__main__':
    app.run(debug=True)
<!doctype html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>WERS Code Extractor</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body>
    <!-- Navbar -->
    <nav class="navbar navbar-expand-lg navbar-dark bg-dark">
        <div class="container-fluid">
            <a class="navbar-brand" href="#">WERS Code Extractor</a>
            <button class="navbar-toggler" type="button" data-bs-toggle="collapse" data-bs-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
                <span class="navbar-toggler-icon"></span>
            </button>
            <div class="collapse navbar-collapse" id="navbarNav">
                <ul class="navbar-nav ms-auto">
                    <li class="nav-item">
                        <a class="nav-link" href="#">Home</a>
                    </li>
                    <li class="nav-item">
                        <a class="nav-link" href="#">Upload</a>
                    </li>
                    <li class="nav-item">
                        <a class="nav-link" href="#">Results</a>
                    </li>
                </ul>
            </div>
        </div>
    </nav>

    <!-- Main Content -->
    <div class="container mt-5">
        <h1 class="text-center">Upload Word Documents and Input Codes</h1>
        <form method="POST" enctype="multipart/form-data" class="mt-4">
            <div class="mb-3">
                <label for="file" class="form-label">Upload Word Document 1:</label>
                <input type="file" name="file" class="form-control" required>
            </div>

            <div class="mb-3">
                <label for="file2" class="form-label">Upload Word Document 2 (Optional):</label>
                <input type="file" name="file2" class="form-control">
            </div>

            <div class="mb-3">
                <label for="input_codes" class="form-label">Enter WERS Codes (one per line):</label>
                <textarea name="input_codes" rows="5" class="form-control" placeholder="Enter WERS Codes"></textarea>
            </div>
            
            <div class="mb-3">
                <label for="voci_codes" class="form-label">Paste VOCI Codes:</label>
                <textarea name="voci_codes" rows="5" class="form-control" placeholder="Paste VOCI Codes"></textarea>
            </div>
            
            <button type="submit" class="btn btn-primary w-100">Upload</button>
        </form>
    </div>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>



<!doctype html>
<html>
<head>
    <title>WERS Code Extractor - Results</title>
    <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
    <style>
        .container {
            margin-top: 20px;
        }
        .voci-only {
            color: red;
        }
        .common-voci-doc {
            color: green;
        }
        .doc2 {
            color: blue;
        }
        .doc1 {
            color: orange;
        }
        .doc1-doc2 {
            color: purple;
        }
        .filter-section {
            margin-bottom: 20px;
        }
    </style>
</head>
<body>
    <nav class="navbar navbar-expand-lg navbar-light bg-light">
        <a class="navbar-brand" href="#">WERS Code Extractor</a>
        <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
            <span class="navbar-toggler-icon"></span>
        </button>
        <div class="collapse navbar-collapse" id="navbarNav">
            <ul class="navbar-nav">
                <li class="nav-item">
                    <a class="nav-link" href="/">Home</a>
                </li>
            </ul>
        </div>
    </nav>

    <div class="container">
        <h1 class="my-4">Processing Results</h1>
        <p><a href="{{ file_txt_url }}" class="btn btn-secondary">Download the text file</a></p>
        
        <div class="filter-section">
            <h4>Filter Results</h4>
            <div class="btn-group" role="group">
                <button type="button" class="btn btn-outline-primary filter-btn active" data-filter="all">All</button>
                <button type="button" class="btn btn-outline-primary filter-btn" data-filter="VOCI Only">VOCI Only</button>
                <button type="button" class="btn btn-outline-primary filter-btn" data-filter="WERS Document 1 Only">WERS Doc 1</button>
                <button type="button" class="btn btn-outline-primary filter-btn" data-filter="WERS Document 2 Only">WERS Doc 2</button>
                <button type="button" class="btn btn-outline-primary filter-btn" data-filter="Both">Both VOCI & WERS</button>
            </div>
        </div>
        
        <h2 class="my-4">Results</h2>
        {% if code_results %}
            <div class="results">
                <table class="table table-striped" id="resultsTable">
                    <thead>
                        <tr>
                            <th>S.No.</th>
                            <th>Code</th>
                            <th>Description</th>
                            <th>Source</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for code, result in code_results %}
                        <tr data-category="{{ result.source }}">
                            <td>{{ loop.index }}</td>
                            <td class="{% if result.source == 'VOCI Only' %}voci-only{% elif result.source == 'Both VOCI and WERS Document 1' %}common-voci-doc{% elif result.source == 'Both VOCI and WERS Document 2' %}common-voci-doc{% elif result.source == 'WERS Document 1 Only' %}doc1{% elif result.source == 'WERS Document 2 Only' %}doc2{% elif result.source == 'Both VOCI and WERS Document 1 and 2' %}common-voci-doc{% endif %}">
                                {{ code }}
                            </td>
                            <td>{{ result.description }}</td>
                            <td>{{ result.source }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        {% else %}
            <p>No codes found.</p>
        {% endif %}
    </div>

    <script src="https://code.jquery.com/jquery-3.5.1.slim.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/@popperjs/core@2.9.1/dist/umd/popper.min.js"></script>
    <script src="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/js/bootstrap.min.js"></script>
    <script>
        $(document).ready(function() {
            $('.filter-btn').click(function() {
                // Remove active class from all buttons
                $('.filter-btn').removeClass('active');
                // Add active class to clicked button
                $(this).addClass('active');
                
                const filter = $(this).data('filter');
                const rows = $('#resultsTable tbody tr');
                let visibleIndex = 1;
                
                rows.each(function() {
                    const category = $(this).data('category');
                    if (filter === 'all') {
                        $(this).show();
                        $(this).find('td:first').text(visibleIndex++);
                    } else if (filter === 'Both') {
                        if (category.includes('Both VOCI and WERS')) {
                            $(this).show();
                            $(this).find('td:first').text(visibleIndex++);
                        } else {
                            $(this).hide();
                        }
                    } else {
                        if (category === filter) {
                            $(this).show();
                            $(this).find('td:first').text(visibleIndex++);
                        } else {
                            $(this).hide();
                        }
                    }
                });
            });
        });
    </script>
</body>
</html>
