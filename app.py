from flask import Flask, request, render_template, send_from_directory, url_for
import re
import os
from docx import Document

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []

    # Extract visible text from paragraphs with section numbers
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():  # Only add non-empty paragraphs
            full_text.append(f"{i}. {para.text}")

    # Extract text from tables
    for table_num, table in enumerate(doc.tables, 1):
        full_text.append(f"\nTable {table_num}:")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))

    # Extract text from headers and footers
    for section_num, section in enumerate(doc.sections, 1):
        # Extract headers
        header = section.header
        if any(paragraph.text.strip() for paragraph in header.paragraphs):
            full_text.append(f"\nHeader Section {section_num}:")
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

        # Extract footers
        footer = section.footer
        if any(paragraph.text.strip() for paragraph in footer.paragraphs):
            full_text.append(f"\nFooter Section {section_num}:")
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

    return '\n'.join(full_text)

def extract_codes_from_text(text, codes):
    codes_set = set(codes)
    found_codes = set()

    # Match each code in the text
    for code in codes_set:
        if code in text:
            found_codes.add(code)

    return found_codes

def extract_codes_and_descriptions(doc):
    codes_with_descriptions = {}

    # First pass: Look for descriptions followed by codes in paragraphs and tables
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Remove any "Note:" sections
        text = re.sub(r'Note:.*$', '', text).strip()

        # Look for description followed by code pattern
        # Match text ending with a 4-5 character alphanumeric code
        match = re.search(r'^(.*?)\s*([A-Z0-9]{4,5})$', text)
        if match:
            desc, code = match.groups()
            desc = desc.strip(' -–')  # Remove leading/trailing dashes
            if desc:
                codes_with_descriptions[code] = desc

    # Second pass: Look in tables
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 1:
                continue

            # Get full row text
            row_text = ' '.join(cell.text.strip() for cell in row.cells)
            row_text = re.sub(r'Note:.*$', '', row_text).strip()

            # Look for description followed by code pattern
            match = re.search(r'^(.*?)\s*([A-Z0-9]{4,5})$', row_text)
            if match:
                desc, code = match.groups()
                desc = desc.strip(' -–')
                if desc:
                    codes_with_descriptions[code] = desc

            # Also check individual cells for the pattern
            for cell in row.cells:
                cell_text = cell.text.strip()
                match = re.search(r'^(.*?)\s*([A-Z0-9]{4,5})$', cell_text)
                if match:
                    desc, code = match.groups()
                    desc = desc.strip(' -–')
                    if desc and code not in codes_with_descriptions:
                        codes_with_descriptions[code] = desc

    return codes_with_descriptions

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files.get('file')
        file2 = request.files.get('file2')
        input_codes = request.form.get('input_codes', '')
        voci_codes_text = request.form.get('voci_codes', '')

        if file:
            # Load the document
            doc = Document(file)

            # Extract descriptions first
            codes_with_descriptions = extract_codes_and_descriptions(doc)

            # Convert DOCX to text for code matching
            text = extract_text_from_docx(file)

            # Parse input WERS codes from newline-separated input while preserving order
            input_codes_list = [code for code in re.findall(r'[A-Z0-9]{5}', input_codes)]

            # Parse VOCI codes from pasted text
            voci_codes_list = re.findall(r'[A-Z0-9]{5}', voci_codes_text)

            # Extract codes found in Word document 1
            found_codes_doc1 = extract_codes_from_text(text, input_codes_list)

            found_codes_doc2 = set()
            if file2:
                # Extract text from second document
                text2 = extract_text_from_docx(file2)
                found_codes_doc2 = extract_codes_from_text(text2, input_codes_list)

            # Determine which codes are from VOCI alone
            voci_alone_codes = set(voci_codes_list) - found_codes_doc1 - found_codes_doc2

            # Determine codes common to both VOCI and Word documents
            common_voci_and_doc1 = set(voci_codes_list).intersection(found_codes_doc1)
            common_voci_and_doc2 = set(voci_codes_list).intersection(found_codes_doc2)

            # Combine all found codes for display
            combined_codes = found_codes_doc1.union(found_codes_doc2)

            # Modify the code_results creation to include descriptions
            code_results = []
            for code in input_codes_list:
                if code in voci_alone_codes:
                    source = 'VOCI Only'
                elif code in common_voci_and_doc1 and code in common_voci_and_doc2:
                    source = 'Both VOCI and WERS Document 1 and 2'
                elif code in common_voci_and_doc1:
                    source = 'Both VOCI and WERS Document 1'
                elif code in common_voci_and_doc2:
                    source = 'Both VOCI and WERS Document 2'
                elif code in found_codes_doc1:
                    source = 'WERS Document 1 Only'
                elif code in found_codes_doc2:
                    source = 'WERS Document 2 Only'
                else:
                    continue

                result_dict = {
                    'source': source,
                    'description': codes_with_descriptions.get(code, '')
                }
                code_results.append((code, result_dict))

            # Add VOCI-only codes with their descriptions
            for code in voci_codes_list:
                if code not in input_codes_list:
                    result_dict = {
                        'source': 'VOCI Only',
                        'description': codes_with_descriptions.get(code, '')
                    }
                    code_results.append((code, result_dict))

            # Save text file with extracted text and results
            text_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'file.txt')

            # Calculate time metrics for CFD completion (excluding VOCI-only codes)
            # Filter out VOCI-only codes
            wers_codes = [code for code, result in code_results if 'VOCI Only' not in result['source']]

            # Always add 1 day buffer time for entity and MPV$ codes
            entity_mpv_codes = [code for code in wers_codes if 'ENTITY' in code.upper() or 'MPV$' in code.upper()]
            buffer_days = 1  # Always add 1 day buffer for entity/MPV$ codes

            total_codes = len(wers_codes)
            total_minutes = total_codes * 4  # Each code takes 4 minutes
            total_hours = total_minutes / 60
            total_days = (total_hours / 8) + buffer_days  # Assuming 8 working hours per day + buffer

            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write("WERS Code Analysis Results\n")
                f.write("=========================\n\n")

                # Write time metrics
                f.write("CFD Completion Time Estimate (WERS Codes Only):\n")
                f.write("-----------------------------------------\n")
                f.write(f"Total WERS Codes (excluding VOCI-only): {total_codes}\n")
                f.write(f"Total Minutes (4 mins per code): {total_minutes}\n")
                f.write(f"Total Hours: {round(total_hours, 2)}\n")
                f.write(f"Total Working Days (8hrs/day + 1 day buffer): {round(total_days, 2)}\n")
                f.write(f"Note: 1 day buffer is added for entity and MPV$ codes\n\n")

                # Write Document 1 content
                f.write("Extracted Text from Document 1:\n")
                f.write("--------------------------\n")
                f.write(text)
                f.write("\n\n")

                # Write Document 2 content if it exists
                if file2:
                    f.write("Extracted Text from Document 2:\n")
                    f.write("--------------------------\n")
                    f.write(text2)
                    f.write("\n\n")

                # Write analysis results
                f.write("Analysis Results:\n")
                f.write("----------------\n")
                for code, result in code_results:
                    f.write(f"{code}: {result['source']} - {result['description']}\n")

            # Calculate base days without buffer
            base_days = total_hours / 8

            return render_template('display.html',
                                   code_results=code_results,
                                   file_txt_url=url_for('download_file', filename='file.txt'),
                                   time_metrics={
                                       'total_codes': total_codes,
                                       'total_minutes': total_minutes,
                                       'total_hours': round(total_hours, 2),
                                       'base_days': round(base_days, 2),
                                       'buffer_days': buffer_days,
                                       'total_days': round(total_days, 2),
                                       'has_entity_mpv': len(entity_mpv_codes) > 0
                                   })

    return render_template('upload.html')

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@app.route('/process', methods=['POST'])
def process():
    # This route is a placeholder and may need to be implemented fully
    # Currently the main functionality is in the upload_file route

    # If this route is used, we should implement the time metrics calculation here too
    return render_template('upload.html')

if __name__ == '__main__':
    app.run(debug=True)
